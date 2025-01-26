from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import logging
from .openai_service import SOPGenerator

def generate_sop_document(sop_data):
    try:
        doc = Document()
        sop_generator = SOPGenerator()

        # Page setup
        sections = doc.sections
        for section in sections:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

        # Basic document setup - Century Gothic 10pt
        style = doc.styles['Normal']
        style.font.name = 'Century Gothic'
        style.font.size = Pt(10)

        # Update all built-in styles to use Century Gothic
        for style_name in ['Heading1', 'Heading2', 'Title']:
            if style_name in doc.styles:
                style = doc.styles[style_name]
                style.font.name = 'Century Gothic'
                style.font.size = Pt(10)
                style.font.bold = True

        # Title
        title = doc.add_heading(sop_data['title'], level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = 'Century Gothic'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.underline = True

        # Add spacing after title
        doc.add_paragraph()

        # Document info table
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        info_table.autofit = True

        # Add document info
        info = [
            ('Document ID:', sop_data['document_id']),
            ('Effective Date:', sop_data['effective_date'].strftime('%m/%d/%Y')),
            ('Version:', sop_data['version']),
            ('Created:', datetime.now().strftime('%m/%d/%Y'))
        ]

        for i, (label, value) in enumerate(info):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
            # Make label cells bold
            row.cells[0].paragraphs[0].runs[0].font.bold = True

        # Add spacing after info table
        doc.add_paragraph()

        # Generate content using OpenAI
        content = sop_generator.generate_sop_content(sop_data)
        if content:
            # Split content into sections and add them to document
            sections = content.split('\n\n')
            for section in sections:
                if section.strip():
                    lines = section.strip().split('\n')
                    # Check if the line is a heading (numbered section)
                    if lines[0].strip()[0].isdigit() and '.' in lines[0]:
                        heading = doc.add_paragraph()
                        run = heading.add_run(lines[0].strip())
                        run.font.name = 'Century Gothic'
                        run.font.size = Pt(10)
                        run.bold = True
                        # Only underline main section headers (1-9)
                        # Do not underline procedure steps (section 5 and its subsections)
                        if len(lines[0].split('.')[0].strip()) == 1 and not lines[0].strip().startswith('5'):
                            run.underline = True
                        # Add remaining lines
                        for line in lines[1:]:
                            para = doc.add_paragraph(line.strip())
                            para.style.font.name = 'Century Gothic'
                            para.style.font.size = Pt(10)
                    else:
                        para = doc.add_paragraph()
                        run = para.add_run(section.strip())
                        run.font.name = 'Century Gothic'
                        run.font.size = Pt(10)

        # Add spacing before contacts
        doc.add_paragraph()

        # Contact Information Section Header
        contacts_heading = doc.add_heading('Contact Information', level=2)
        contacts_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in contacts_heading.runs:
            run.font.name = 'Century Gothic'
            run.font.size = Pt(10)
            run.bold = True
            run.underline = True

        # IT Delivery Team Section
        contacts = doc.add_paragraph()
        team_header = contacts.add_run('HC IT Delivery Team')
        team_header.font.name = 'Century Gothic'
        team_header.font.size = Pt(10)
        team_header.bold = True
        team_header.underline = True

        # Add expand/collapse hint for IT section
        hint_para = doc.add_paragraph()
        hint_run = hint_para.add_run('(Click + to expand or - to collapse contact details)')
        hint_run.font.name = 'Century Gothic'
        hint_run.font.size = Pt(10)
        hint_run.italic = True

        # Add IT contact information
        contacts = doc.add_paragraph()
        contacts.add_run('+ Contact 1: ').bold = True
        contacts.add_run('Veronica Nolte\n')
        contacts.add_run('    Role: _______________________\n')
        contacts.add_run('    Email: \n')
        contacts.add_run('    Phone: (021) 111-111\n\n')

        contacts.add_run('+ Contact 2: ').bold = True
        contacts.add_run('XXXXX\n')
        contacts.add_run('    Role: _______________________\n')
        contacts.add_run('    Email: \n')
        contacts.add_run('    Phone: (021) 111-112\n\n')

        contacts.add_run('+ Contact 3: ').bold = True
        contacts.add_run('XXXXX\n')
        contacts.add_run('    Role: _______________________\n')
        contacts.add_run('    Email: \n')
        contacts.add_run('    Phone: (021) 111-113\n\n')

        # Payroll Support Section
        payroll_header = doc.add_paragraph()
        header_run = payroll_header.add_run('HCSC Payroll Support')
        header_run.font.name = 'Century Gothic'
        header_run.font.size = Pt(10)
        header_run.bold = True
        header_run.underline = True

        # Add expand/collapse hint for payroll
        hint_para = doc.add_paragraph()
        hint_run = hint_para.add_run('(Click + to expand or - to collapse contact details)')
        hint_run.font.name = 'Century Gothic'
        hint_run.font.size = Pt(10)
        hint_run.italic = True

        # Add Payroll contact template
        payroll_contacts = doc.add_paragraph()
        for i in range(1, 4):
            payroll_contacts.add_run(f'+ Contact {i}: ').bold = True
            payroll_contacts.add_run('_______________________\n')
            payroll_contacts.add_run('    Role: _______________________\n')
            payroll_contacts.add_run('    Email: _______________________\n')
            payroll_contacts.add_run('    Phone: _______________________\n\n')

        # Authorization Section
        auth_heading = doc.add_heading('AUTHORIZED BY:', level=2)
        auth_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in auth_heading.runs:
            run.font.name = 'Century Gothic'
            run.font.size = Pt(10)
            run.bold = True
            run.underline = True

        # Process Owner
        auth = doc.add_paragraph()
        auth.add_run('Process Owner:\n').bold = True
        auth.add_run('Name & Surname: _______________________\n')
        auth.add_run('Role: _______________________\n')
        auth.add_run('Signature & date: _______________________\n\n')

        # Area Head
        auth.add_run('Area Head:\n').bold = True
        auth.add_run('Name & Surname: _______________________\n')
        auth.add_run('Role: _______________________\n')
        auth.add_run('Signature & date: _______________________')

        # Ensure consistent font throughout the document
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Century Gothic'
                run.font.size = Pt(10)

        return doc

    except Exception as e:
        logging.error(f"Error generating SOP document: {str(e)}")
        raise